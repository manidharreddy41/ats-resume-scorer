#!/usr/bin/env python3
"""
ATS Resume Scorer — FastAPI Backend
Uses Claude API for intelligent ATS scoring (mirrors real Workday/Greenhouse/Lever logic).
Preserves exact DOCX formatting on output; PDF→DOCX for editing.
"""

import os, re, json, uuid, shutil, copy
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
import subprocess

BASE    = Path(__file__).parent
UPLOADS = BASE / "uploads"
OUTPUTS = BASE / "outputs"
STATIC  = BASE / "static"

UPLOADS.mkdir(exist_ok=True)
OUTPUTS.mkdir(exist_ok=True)

# ── Claude runner: tries CLI first (no key needed), falls back to SDK key ─────
def _find_claude_cli() -> str:
    """Find the claude CLI binary."""
    for candidate in ["claude", r"C:\Users\nagam\AppData\Local\AnthropicClaude\claude.exe",
                      r"C:\Program Files\Claude\claude.exe"]:
        try:
            result = subprocess.run([candidate, "--version"], capture_output=True, timeout=5)
            if result.returncode == 0:
                return candidate
        except Exception:
            pass
    return ""

def _load_env_key() -> str:
    for env_file in [BASE / ".env", Path("D:/claude/.env"), Path("D:/claude/GameBudget-API/.env")]:
        try:
            for line in env_file.read_text().splitlines():
                if line.startswith("ANTHROPIC_API_KEY="):
                    k = line.split("=", 1)[1].strip().strip('"').strip("'")
                    if k:
                        return k
        except Exception:
            pass
    return ""

_CLAUDE_CLI     = _find_claude_cli()
_SERVER_SDK_KEY = os.environ.get("ANTHROPIC_API_KEY", "") or _load_env_key()

# Server has native auth if CLI works OR SDK key is set
SERVER_HAS_KEY = bool(_CLAUDE_CLI or _SERVER_SDK_KEY)

def call_claude(prompt: str, request_key: str = "") -> str:
    """Call Claude via CLI (preferred — uses existing login) or SDK as fallback."""
    import time
    # ── CLI path (uses your Claude Code login, no key needed) ────────────────
    if _CLAUDE_CLI:
        for attempt in range(3):
            try:
                # Pass as raw bytes to avoid Windows cp1252 encoding issues with
                # special chars (em dashes, smart quotes) extracted from PDFs/DOCX.
                result = subprocess.run(
                    [_CLAUDE_CLI, "-p", "--input-format", "text"],
                    input=prompt.encode("utf-8"),
                    capture_output=True, timeout=180,
                    env=os.environ.copy()
                )
                stdout = result.stdout.decode("utf-8", errors="replace").strip()
                stderr = result.stderr.decode("utf-8", errors="replace").strip()
                if result.returncode == 0 and stdout:
                    return stdout
                print(f"[claude CLI] attempt={attempt+1} rc={result.returncode} stderr={stderr[:300]}")
                if attempt < 2:
                    time.sleep(3)  # wait before retry (rate limit / transient failure)
            except subprocess.TimeoutExpired:
                print(f"[claude CLI] attempt={attempt+1} timed out")
                if attempt < 2:
                    time.sleep(2)
            except Exception as e:
                print(f"[claude CLI] attempt={attempt+1} exception: {e}")
                if attempt < 2:
                    time.sleep(2)

    # ── SDK path (uses server key or user-supplied key) ───────────────────────
    import anthropic
    key = _SERVER_SDK_KEY or request_key
    if not key:
        raise HTTPException(503, "Claude CLI unavailable and no API key configured. Please wait a moment and try again, or enter your Anthropic API key.")
    client = anthropic.Anthropic(api_key=key)
    msg = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}]
    )
    return msg.content[0].text.strip()


def extract_json(text: str) -> str:
    """Extract the first complete JSON object from text that may have extra content."""
    # Strip markdown fences
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    text = text.strip()
    # Find the first { and match its closing }
    start = text.find('{')
    if start == -1:
        return text
    depth = 0
    in_str = False
    escape = False
    for i, ch in enumerate(text[start:], start):
        if escape:
            escape = False
            continue
        if ch == '\\' and in_str:
            escape = True
            continue
        if ch == '"':
            in_str = not in_str
            continue
        if not in_str:
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    return text[start:i+1]
    return text[start:]

# ── Resume text extraction ────────────────────────────────────────────────────
def extract_text_docx(path: Path) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text_pdf(path: Path) -> str:
    with pdfplumber.open(path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

# ── JD text from URL ──────────────────────────────────────────────────────────
def fetch_jd_from_url(url: str) -> str:
    try:
        r = requests.get(url, timeout=12, headers={"User-Agent": "Mozilla/5.0"})
        from html.parser import HTMLParser
        class _P(HTMLParser):
            def __init__(self):
                super().__init__(); self.chunks = []; self._skip = False
            def handle_starttag(self, t, a):
                if t in ("script","style","nav","header","footer","aside"): self._skip = True
            def handle_endtag(self, t):
                if t in ("script","style","nav","header","footer","aside"): self._skip = False
            def handle_data(self, d):
                if not self._skip and d.strip(): self.chunks.append(d.strip())
        p = _P(); p.feed(r.text)
        return " ".join(p.chunks)[:8000]
    except Exception:
        return ""

# ── Claude ATS scoring ────────────────────────────────────────────────────────
ATS_PROMPT = """You are an expert ATS (Applicant Tracking System) analyst trained on Workday, Greenhouse, Lever, and iCIMS scoring logic. Analyze this resume against the job description exactly as a real ATS + recruiter AI would.

SCORING CRITERIA (total 100 points):
- Keyword Coverage (30 pts): Exact and semantic matches for required/preferred skills in JD
- Technical Skills Alignment (25 pts): Depth of match for tools, languages, frameworks
- Experience Relevance (20 pts): Job titles, responsibilities, domain overlap
- Quantified Impact (15 pts): Metrics, numbers, business outcomes present
- ATS Formatting (10 pts): Proper section headers, no tables/graphics hiding text, parseable structure

For experience bullets, you MUST identify which specific job/company in the resume the bullet should be added under. Use the exact company name and job title as they appear in the resume (e.g. "AI/ML Engineer — Apple Inc." or "Data Analyst — T-Mobile").

Respond ONLY with a valid JSON object matching this exact schema:

{
  "score": <integer 0-100>,
  "grade": "<Excellent|Good|Fair|Needs Work>",
  "breakdown": {
    "keyword_coverage": <0-30>,
    "technical_skills": <0-25>,
    "experience_relevance": <0-20>,
    "quantified_impact": <0-15>,
    "ats_formatting": <0-10>
  },
  "matched_keywords": ["<keyword>"],
  "missing_keywords": [
    {"keyword": "<keyword>", "importance": "<Required|Preferred>", "pts_impact": <integer>}
  ],
  "suggestions": [
    {
      "id": "<unique_id>",
      "section": "<Technical Skills|Experience|Summary>",
      "action": "<add_to_skills|add_bullet|update_summary>",
      "text": "<exact sentence or bullet to add — write it ready to paste, professional tone, first person implied>",
      "target_role": "<for Experience bullets: copy the EXACT employer name from the resume's experience section, e.g. 'ASPYRANT LLC' or 'AETNA'. Use the most recent employer if you're unsure. Empty string for Skills/Summary suggestions only>",
      "keywords_targeted": ["<kw>"],
      "reason": "<why this specifically helps pass ATS and impress recruiters>",
      "pts_gain": <estimated score increase as integer>
    }
  ],
  "ats_tips": ["<actionable formatting or content tip>"],
  "summary": "<2-3 sentence honest assessment of interview likelihood with this resume for this role>"
}

JOB DESCRIPTION:
{JD}

RESUME:
{RESUME}"""

def run_ats_analysis(resume_text: str, jd_text: str, request_key: str = "") -> dict:
    prompt = ATS_PROMPT.replace("{JD}", jd_text[:5000]).replace("{RESUME}", resume_text[:8000])
    raw = call_claude(prompt, request_key)
    return json.loads(extract_json(raw))

# ── Re-score after applying suggestions ───────────────────────────────────────
RESCORE_PROMPT = """You are an ATS scoring system. The candidate applied suggested improvements to their resume. Re-score it against the same job description.

Respond ONLY with valid JSON:
{
  "score": <integer 0-100>,
  "grade": "<Excellent|Good|Fair|Needs Work>",
  "matched_keywords": ["<keyword>"],
  "missing_keywords": [{"keyword": "<keyword>", "importance": "<Required|Preferred>"}],
  "summary": "<1-2 sentences on interview likelihood now>"
}

JOB DESCRIPTION:
{JD}

UPDATED RESUME:
{RESUME}"""

def run_rescore(resume_text: str, jd_text: str, request_key: str = "") -> dict:
    prompt = RESCORE_PROMPT.replace("{JD}", jd_text[:5000]).replace("{RESUME}", resume_text[:8000])
    raw = call_claude(prompt, request_key)
    return json.loads(extract_json(raw))

# ── Document structure helpers ────────────────────────────────────────────────
_BULLET_CHARS = ('•', '–', '○', '·', '▪', '◦', '‣')
_SECTION_RE   = re.compile(
    r'^(professional\s+)?(work\s+)?'
    r'(experience|employment|history|education|skills|technical\s+skills|'
    r'summary|professional\s+summary|objective|profile|projects|certifications)'
    r'\s*:?\s*$',
    re.I
)
_JOB_SEP_RE  = re.compile(r'[|@]')
_DATE_RE     = re.compile(
    r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|march|'
    r'april|june|july|august|september|october|november|december)\w*[\s,]+\d{4}'
    r'|\b\d{4}\s*[-–—]\s*(\d{4}|present|current|till\s+date|to\s+date)\b', re.I)

def _is_bullet(para) -> bool:
    t = para.text.strip()
    return t[:1] in _BULLET_CHARS

def _is_section_header(para) -> bool:
    t = para.text.strip()
    return bool(t and _SECTION_RE.match(t) and len(t) < 80)

def _is_job_header(para) -> bool:
    """Company/role line: has a date range (most reliable signal for a job header)."""
    t = para.text.strip()
    if not t or len(t) > 250 or t[:1] in _BULLET_CHARS:
        return False
    return bool(_DATE_RE.search(t))

def _global_bullet_template(paras):
    """
    Return the most representative bullet paragraph for use as a formatting template.
    Finds the most common indent level among bullets and returns a bullet with that indent.
    """
    bullets = [p for p in paras if _is_bullet(p) and p.text.strip()]
    if not bullets:
        return None
    # Find modal indent (most common among bullets)
    from collections import Counter
    indent_counts = Counter(
        (p.paragraph_format.left_indent or 0) for p in bullets
    )
    modal_indent = indent_counts.most_common(1)[0][0]
    # Return first bullet with that indent
    for p in bullets:
        if (p.paragraph_format.left_indent or 0) == modal_indent:
            return p
    return bullets[0]

def _find_job_sections(paras):
    """
    Scan paragraphs and return list of:
      {"header": str, "header_idx": int, "bullet_end_idx": int}
    where bullet_end_idx is the index AFTER the last bullet of this section.
    """
    sections = []
    in_exp = False

    for i, p in enumerate(paras):
        t = p.text.strip()
        if not t:
            continue
        if _is_section_header(p):
            in_exp = bool(re.search(r'\bexperience\b', t, re.I))
            continue
        if in_exp and _is_job_header(p):
            # only first line of multi-line paras
            header_line = t.split('\n')[0].strip()
            sections.append({"header": header_line, "header_idx": i, "bullet_end_idx": i + 1})

    # For each section, walk forward to find last bullet index
    for k, sec in enumerate(sections):
        next_start = sections[k + 1]["header_idx"] if k + 1 < len(sections) else len(paras)
        last_bullet_idx = sec["header_idx"]
        for j in range(sec["header_idx"] + 1, next_start):
            t = paras[j].text.strip()
            if not t:
                continue
            ind = paras[j].paragraph_format.left_indent or 0
            # Count as a content line: has bullet char, has positive indent, or is a
            # plain sentence (not a section/job header) within the experience block.
            if (_is_bullet(paras[j]) or ind > 0 or
                    (not _is_section_header(paras[j]) and not _is_job_header(paras[j]))):
                last_bullet_idx = j
        sec["bullet_end_idx"] = last_bullet_idx  # index of last bullet (insert AFTER this)

    return sections

def _match_job_section(sections, target_role: str):
    """Fuzzy match target_role to the best section header."""
    if not target_role or not sections:
        return None
    tl = target_role.lower()
    words = [w for w in re.split(r'\W+', tl) if len(w) > 2]
    best, best_score = None, 0
    for sec in sections:
        hl = sec["header"].lower()
        score = sum(1 for w in words if w in hl)
        if score > best_score:
            best, best_score = sec, score
    return best if best_score > 0 else (sections[0] if sections else None)

def _make_bullet_element(text: str, template_para, prefix: str = '•'):
    """
    Build a new <w:p> element for a bullet, copying paragraph + run properties
    from template_para (a real bullet paragraph). Adds the bullet prefix char.
    """
    tmpl = template_para._element
    new_p = copy.deepcopy(tmpl)

    # Strip all existing runs from the copy
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)

    # Build new run: copy rPr from first run of template
    r_elem = OxmlElement("w:r")
    orig_runs = tmpl.findall(qn("w:r"))
    if orig_runs:
        rpr = orig_runs[0].find(qn("w:rPr"))
        if rpr is not None:
            new_rpr = copy.deepcopy(rpr)
            # Remove explicit bold override so it inherits naturally
            for tag in (qn("w:b"), qn("w:bCs")):
                el = new_rpr.find(tag)
                if el is not None:
                    new_rpr.remove(el)
            r_elem.append(new_rpr)

    t_elem = OxmlElement("w:t")
    # Keep the same bullet prefix the rest of the doc uses
    existing_prefix = template_para.text.strip()[:1]
    bullet_prefix   = existing_prefix if existing_prefix in _BULLET_CHARS else prefix
    t_elem.text = bullet_prefix + " " + text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    new_p.append(r_elem)
    return new_p

def _make_plain_element(text: str, template_para):
    """Build a plain <w:p> element copying pPr+rPr from template_para (no bullet prefix)."""
    tmpl = template_para._element
    new_p = copy.deepcopy(tmpl)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    r_elem = OxmlElement("w:r")
    orig_runs = tmpl.findall(qn("w:r"))
    if orig_runs:
        rpr = orig_runs[0].find(qn("w:rPr"))
        if rpr is not None:
            new_rpr = copy.deepcopy(rpr)
            for tag in (qn("w:b"), qn("w:bCs")):
                el = new_rpr.find(tag)
                if el is not None:
                    new_rpr.remove(el)
            r_elem.append(new_rpr)
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    new_p.append(r_elem)
    return new_p

# ── Apply suggestions to DOCX preserving all formatting ──────────────────────
def apply_to_docx(source_path: Path, selected: list[dict]) -> Path:
    doc    = Document(source_path)
    paras  = doc.paragraphs
    sections = _find_job_sections(paras)
    # Get a real bullet paragraph as the formatting template for new bullets
    bullet_tmpl = _global_bullet_template(paras)

    # ── 1. Skills additions ────────────────────────────────────────────────
    skills_to_add = [kw for s in selected
                     if s["action"] in ("add_to_skills",) or
                        (s["action"] == "add_bullet" and s.get("section") == "Technical Skills")
                     for kw in s.get("keywords_targeted", [])]

    if skills_to_add:
        for i, para in enumerate(paras):
            if any(m in para.text.lower()
                   for m in ["technical skills", "skills", "technologies", "tools", "tech stack"]):
                for j in range(i + 1, min(i + 8, len(paras))):
                    sp = paras[j]
                    if sp.text.strip():
                        existing = sp.text.rstrip()
                        sep = " | " if "|" in existing else ", "
                        new_kws = sep.join(k for k in skills_to_add
                                           if k.lower() not in existing.lower())
                        if new_kws:
                            if sp.runs:
                                lr = sp.runs[-1]
                                nr = sp.add_run(sep + new_kws)
                                nr.bold      = lr.bold
                                nr.font.size = lr.font.size
                                nr.font.name = lr.font.name
                            else:
                                sp.add_run(sep + new_kws)
                        break
                break

    # ── 2. Experience bullets — placed under the correct company ──────────
    exp_bullets = [(s["text"], s.get("target_role", "")) for s in selected
                   if s["action"] == "add_bullet" and s.get("section") == "Experience"]

    if exp_bullets:
        for bullet_text, target_role in exp_bullets:
            sec = _match_job_section(sections, target_role)
            if sec:
                insert_after = paras[sec["bullet_end_idx"]]
            elif sections:
                insert_after = paras[sections[0]["bullet_end_idx"]]
            else:
                insert_after = None

            if insert_after:
                if bullet_tmpl:
                    new_p = _make_bullet_element(bullet_text, bullet_tmpl)
                else:
                    # No bullet chars in resume — copy the style of the insert_after paragraph
                    new_p = _make_plain_element(bullet_text, insert_after)
                insert_after._element.addnext(new_p)

    # ── 3. Summary updates ─────────────────────────────────────────────────
    summary_updates = [s["text"] for s in selected if s["action"] == "update_summary"]
    if summary_updates:
        for i, para in enumerate(paras):
            if (re.search(r'\b(summary|objective|profile|about)\b', para.text, re.I)
                    and len(para.text.strip()) < 40):
                for j in range(i + 1, min(i + 4, len(paras))):
                    sp = paras[j]
                    if sp.text.strip():
                        if sp.runs:
                            sp.runs[-1].text += " " + " ".join(summary_updates)
                        break
                break

    out_name = f"tailored_{source_path.stem}_{uuid.uuid4().hex[:6]}.docx"
    out_path  = OUTPUTS / out_name
    doc.save(out_path)
    return out_path

def apply_to_pdf(source_path: Path, selected: list[dict]) -> Path:
    from pdf2docx import Converter
    tmp = OUTPUTS / f"_tmp_{uuid.uuid4().hex[:6]}.docx"
    cv = Converter(str(source_path))
    cv.convert(str(tmp))
    cv.close()
    result = apply_to_docx(tmp, selected)
    tmp.unlink(missing_ok=True)
    return result

# ── FastAPI ───────────────────────────────────────────────────────────────────
app = FastAPI(title="ATS Resume Scorer")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

sessions: dict[str, dict] = {}

@app.get("/api/config")
async def config():
    return JSONResponse({"server_has_key": SERVER_HAS_KEY, "cli": _CLAUDE_CLI})

@app.get("/api/test-claude")
async def test_claude():
    try:
        out = call_claude("Reply with only the word: WORKING")
        return JSONResponse({"status": "ok", "response": out})
    except Exception as e:
        return JSONResponse({"status": "error", "detail": str(e)}, status_code=500)

@app.get("/api/debug-cli")
async def debug_cli():
    """Raw CLI debug — shows exactly what happens with a full-size prompt."""
    import pdfplumber
    test_prompt = ATS_PROMPT.replace("{JD}", "Python Machine Learning Engineer TensorFlow required")
    # Use a short fake resume so we can isolate the issue
    test_prompt = test_prompt.replace("{RESUME}", "Software Engineer with Python experience. Built ML models.")
    result = subprocess.run(
        [_CLAUDE_CLI, "-p", "--input-format", "text"],
        input=test_prompt, capture_output=True, text=True, timeout=120,
        env=os.environ.copy(), encoding="utf-8"
    )
    return JSONResponse({
        "cli": _CLAUDE_CLI,
        "prompt_len": len(test_prompt),
        "returncode": result.returncode,
        "stdout_len": len(result.stdout),
        "stdout_snippet": result.stdout[:300],
        "stderr_snippet": result.stderr[:300],
    })

@app.post("/api/analyze")
async def analyze(
    resume: UploadFile = File(...),
    jd_text: str = Form(""),
    jd_url:  str = Form(""),
    api_key: str = Form(""),
):
    ext = Path(resume.filename).suffix.lower()
    if ext not in (".pdf", ".docx"):
        raise HTTPException(400, "Only PDF and DOCX supported.")

    sid = uuid.uuid4().hex
    upload_path = UPLOADS / f"{sid}{ext}"
    with open(upload_path, "wb") as f:
        shutil.copyfileobj(resume.file, f)

    if not jd_text.strip() and jd_url.strip():
        jd_text = fetch_jd_from_url(jd_url.strip())

    if not jd_text.strip():
        raise HTTPException(400, "Please paste a job description or enter a valid URL.")

    resume_text = extract_text_docx(upload_path) if ext == ".docx" else extract_text_pdf(upload_path)

    result = run_ats_analysis(resume_text, jd_text, api_key)

    sessions[sid] = {
        "api_key": api_key,  # user-supplied key (empty if server has key)
        "upload_path": str(upload_path),
        "ext": ext,
        "resume_text": resume_text,
        "jd_text": jd_text,
        "initial_score": result.get("score", 0),
        "suggestions": result.get("suggestions", []),
        "output_path": None,
        "filename": resume.filename,
    }

    return JSONResponse({"session_id": sid, "filename": resume.filename, **result})


@app.post("/api/apply")
async def apply_suggestions_endpoint(body: dict):
    sid          = body.get("session_id")
    selected_ids = body.get("selected_ids", [])

    if sid not in sessions:
        raise HTTPException(404, "Session expired. Please re-upload your resume.")

    sess = sessions[sid]
    selected = [s for s in sess["suggestions"] if s["id"] in selected_ids]
    if not selected:
        raise HTTPException(400, "No suggestions selected.")

    upload_path = Path(sess["upload_path"])
    ext = sess["ext"]

    out_path = apply_to_docx(upload_path, selected) if ext == ".docx" \
               else apply_to_pdf(upload_path, selected)

    new_text = extract_text_docx(out_path) if out_path.suffix == ".docx" \
               else extract_text_pdf(out_path)

    rescore = run_rescore(new_text, sess["jd_text"], sess.get("api_key", ""))
    sess["output_path"] = str(out_path)

    return JSONResponse({
        "output_file": out_path.name,
        "score": rescore.get("score", 0),
        "grade": rescore.get("grade", ""),
        "matched_keywords": rescore.get("matched_keywords", []),
        "missing_keywords": rescore.get("missing_keywords", []),
        "improvement": rescore.get("score", 0) - sess["initial_score"],
        "summary": rescore.get("summary", ""),
    })


@app.get("/api/download/{filename}")
async def download(filename: str):
    path = OUTPUTS / filename
    if not path.exists():
        raise HTTPException(404, "File not found.")
    media = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"
             if filename.endswith(".docx") else "application/pdf")
    return FileResponse(str(path), media_type=media, filename=filename)


app.mount("/", StaticFiles(directory=str(STATIC), html=True), name="static")
